import os
import re
import time

from bardapi import Bard
from docx.api import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt


def process_document(docx_file):
    input_document = Document(docx_file)
    output_document = Document()

    # Styles for the output document
    output_document.styles["Normal"].font.name = "Times New Roman"
    output_document.styles["Normal"].font.size = Pt(14)
    output_document.styles["Normal"].font.bold = True

    # Create a new style for 'BardAnswer'
    output_document.styles.add_style("BardAnswer", WD_STYLE_TYPE.PARAGRAPH)
    output_document.styles["BardAnswer"].font.name = "Times New Roman"
    output_document.styles["BardAnswer"].font.size = Pt(14)

    # Set up Bard API key
    os.environ[
        "_BARD_API_KEY"
    ] = "eggr84ncQLjMlWaowPulG9FppqKTWVO2uFfRD6GmI9GdJIZWwTc2cA_jiI9j7yO9wJH67g."

    # additional question
    additional_question = " Вплив війни між Україною та Росією"

    # Initialize the last successfully processed paragraph index
    last_successful_index = 0

    # Loop through paragraphs in the input document
    for i, paragraph in enumerate(input_document.paragraphs[last_successful_index:], 1):
        # Get the text of the paragraph
        paragraph_text = paragraph.text + additional_question
        current_index = i + last_successful_index

        print(current_index)
        print(paragraph_text)
        if not any(word.isalpha() for word in paragraph_text.split()):
            print(f"No words in paragraph {current_index}. Stopping iteration.")
            break
        try:
            # Get Bard answer for the paragraph
            bard_answer = Bard().get_answer(paragraph_text)["content"]
            time.sleep(30)
            print(bard_answer)
            # if "[Image of]" in bard_answer:
            if len(bard_answer) > 0:
                # Add the input paragraph with 'Normal' style
                p_input = output_document.add_paragraph(
                    paragraph_text, style="ListNumber"
                )

                # Add the Bard answer with the custom 'BardAnswer' style
                p_answer = output_document.add_paragraph(
                    bard_answer, style="BardAnswer"
                )
                output_document.save("output_iteration.docx")
        except Exception as e:
            # Handle exceptions, e.g., log the error
            print(f"Error processing paragraph {current_index}: {str(e)}")

            # Save the last successful index
            last_successful_index = current_index

    # Save the output document
    output_path = "output.docx"
    output_document.save(output_path)
    print("Everything went ok!")

    return {
        "status": "success",
        "message": "Document processed successfully",
        "document_path": output_path,
    }
